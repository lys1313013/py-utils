import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


class MarkdownToDocxConverter:
    def __init__(self):
        self.document = Document()
        # 设置中文字体
        self.set_chinese_font()

    def set_chinese_font(self):
        """设置中文字体支持"""
        try:
            # 设置全局字体
            self.document.styles['Normal'].font.name = '宋体'
            self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception as e:
            print(f"字体设置警告: {e}")

    def parse_markdown(self, md_content):
        """解析Markdown内容"""
        lines = md_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # 检查标题
            if self.is_heading(line):
                self.add_heading(line)

            # 检查表格
            elif self.is_table_start(lines, i):
                table_lines = self.extract_table(lines, i)
                self.add_table(table_lines)
                i += len(table_lines) - 1

            # 普通段落
            else:
                self.add_paragraph(line)

            i += 1

    def is_heading(self, line):
        """检查是否为标题"""
        return re.match(r'^#{1,6}\s', line)

    def add_heading(self, line):
        """添加标题到Word文档"""
        match = re.match(r'^(#+)\s*(.*)', line)
        if match:
            level = len(match.group(1))  # #的数量决定标题级别
            text = match.group(2).strip()

            # 限制标题级别在1-6之间
            level = max(1, min(level, 6))

            # 添加标题
            heading = self.document.add_heading(text, level=level - 1)

            # 设置中文字体
            try:
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            except:
                pass

    def is_table_start(self, lines, index):
        """检查是否以表格开始"""
        if index + 1 >= len(lines):
            return False

        current_line = lines[index].strip()
        next_line = lines[index + 1].strip()

        # 检查表格分隔线
        if '|' in current_line and '|' in next_line:
            # 下一行应该包含 --- 或 :-- 等表格分隔符
            return re.search(r'[\-:\|]+\s*[\-:\|]*', next_line)
        return False

    def extract_table(self, lines, start_index):
        """提取表格内容"""
        table_lines = []
        i = start_index

        while i < len(lines):
            line = lines[i].strip()
            if '|' in line:
                table_lines.append(line)
                i += 1
            else:
                break

        return table_lines

    def add_table(self, table_lines):
        """添加表格到Word文档"""
        if len(table_lines) < 2:
            return

        # 解析表头
        headers = self.parse_table_row(table_lines[0])

        # 创建表格
        num_cols = len(headers)
        table = self.document.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'

        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header.strip()
            # 设置表头样式
            try:
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
            except:
                pass

        # 添加表格数据
        for i in range(2, len(table_lines)):
            row_data = self.parse_table_row(table_lines[i])
            if len(row_data) == num_cols:
                row_cells = table.add_row().cells
                for j, cell_data in enumerate(row_data):
                    row_cells[j].text = cell_data.strip()

    def parse_table_row(self, row_line):
        """解析表格行"""
        # 移除行首尾的 |
        cleaned_line = row_line.strip()
        if cleaned_line.startswith('|'):
            cleaned_line = cleaned_line[1:]
        if cleaned_line.endswith('|'):
            cleaned_line = cleaned_line[:-1]

        # 分割单元格
        cells = [cell.strip() for cell in cleaned_line.split('|')]
        return cells

    def add_paragraph(self, text):
        """添加普通段落"""
        if text.strip():
            paragraph = self.document.add_paragraph(text.strip())

            # 设置中文字体
            try:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except:
                pass

    def convert_file(self, md_file_path, docx_file_path=None):
        """转换Markdown文件为Word文档"""
        try:
            # 读取Markdown文件
            with open(md_file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # 解析并转换
            self.parse_markdown(md_content)

            # 生成输出文件名
            if docx_file_path is None:
                base_name = os.path.splitext(md_file_path)[0]
                docx_file_path = base_name + '.docx'

            # 保存Word文档
            self.document.save(docx_file_path)
            print(f"转换成功: {md_file_path} -> {docx_file_path}")

            return True

        except Exception as e:
            print(f"转换失败: {e}")
            return False

    def convert_text(self, md_text, docx_file_path):
        """转换Markdown文本为Word文档"""
        try:
            self.parse_markdown(md_text)
            self.document.save(docx_file_path)
            print(f"转换成功: 文本内容 -> {docx_file_path}")
            return True
        except Exception as e:
            print(f"转换失败: {e}")
            return False


def main():
    """主函数 - 使用示例"""
    converter = MarkdownToDocxConverter()

    # 示例1: 转换文件
    md_file = "database_schema.md"  # 替换为您的Markdown文件路径
    converter.convert_file(md_file)


if __name__ == "__main__":
    main()